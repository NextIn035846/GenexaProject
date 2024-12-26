import streamlit as st
from src.GenAI import user_input

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
import base64
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
import io

def create_download_link(val, filename):
    b64 = base64.b64encode(val)  # val looks like b'...'
    return f'<a href="data:application/octet-stream;base64,{b64.decode()}" download="{filename}.docx">Download file</a>'

def set_margins(document, top, bottom, left, right):
    """
    Set margins for the entire document.
    """
    sections = document.sections
    for section in sections:
        section.top_margin = top
        section.bottom_margin = bottom
        section.left_margin = left
        section.right_margin = right

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

def add_header(document, header_text, company_name, logo_path):
    section = document.sections[0]
    header = section.header

    # Add title on the left side
    header_left = header.add_paragraph()
    run_left = header_left.add_run(header_text)
    run_left.bold = True
    run_left.font.size = Pt(10)
    run_left.font.name = 'Arial'

    # Add company name on the right side
    header_right = header.add_paragraph()
    header_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Add company name
    run_company = header_right.add_run(company_name)
    run_company.bold = True
    run_company.font.size = Pt(10)
    run_company.font.name = 'Arial'

    # Add logo
    if logo_path:
        header_right.add_run().add_picture(logo_path, width=Inches(0.5), height=Inches(0.5))

def add_title(document, title_text):
    document.add_heading(title_text, level=1)

    # paragraph = document.add_paragraph()
    # run = paragraph.add_run(title_text)
    # run.bold = True
    # run.font.size = Pt(16)
    # run.font.name = 'Arial'
    # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # document.add_paragraph()  # Add an empty paragraph for spacing
    # document.add_page_break()

def add_paragraph(document, paragraph_text):
    paragraph = document.add_paragraph(paragraph_text)
    paragraph_format = paragraph.paragraph_format
    # paragraph_format.left_indent = Inches(0.5)  # Set left margin
    # paragraph_format.right_indent = Inches(0.5)  # Set right margin
    paragraph_format.space_before = Pt(12)  # Set space before paragraph
    paragraph_format.space_after = Pt(12)  # Set space after paragraph
    run = paragraph.runs[0]
    run.font.size = Pt(12)  # Set font size
    run.font.name = 'Times New Roman'  # Set font name
    document.add_page_break()

def create_document():
    # Create a new Document
    doc = Document()
    # Add Cover Page
    cover_page = doc.add_paragraph("Cover Page", style='Title')
    cover_page.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

    
    # Add header to each page
    header_text = "TITLE OF THE RFP RESPONSE"
    company_name = "Genexa.AI"
    logo_path = "T:\Project Files\pages\Genexa.ai New Logo PNG-1.png"  # Change this to the path of your company logo
    add_header(doc, header_text, company_name, logo_path)

    
    # Add title and paragraph to each page
    add_title(doc,"Cover Letter")
    add_paragraph(doc," ")
    add_title(doc,"Company Background")
    add_paragraph(doc," <Give a brief overview of your organization’s involvement in providing IT services in the marketplace >")
    add_title(doc,"Summary")
    add_paragraph(doc,"Provide Summary of RFP Response in 3 to 4 Paragraphs")
    add_title(doc,"Proposed solution and approach")
    add_paragraph(doc,"""Provide the proposed solution and approach covering the following
a. Key activities and deliverables
b. Timing
c. Information/resource requirements from Family League
d. Deliverables
e. Key milestones, checkpoints, and other decision points""")
    add_title(doc,"Financial Proposal")
    add_paragraph(doc,"Describe the pricing model(s) that you typically employ for your standard services. ")
    
    add_title(doc,"Support")
    add_paragraph(doc,"Describe fully your technical support options, including the assistance request process, escalation process, support hours, response times (for emergency and non-emergency support requests), staffing levels, staff expertise, and physical location of the help desk.")
    add_title(doc,"Case studies and success stories")
    add_paragraph(doc,"Share the information on  companies past performance with the success we have achieved for other clients in the past ")
    
    
    # titles = ["Title 1", "Title 2", "Title 3"]  # Example titles
    # paragraphs = ["Paragraph 1", "Paragraph 2", "Paragraph 3"]  # Example paragraphs
    # for title, paragraph in zip(titles, paragraphs):
    #     add_title(doc, title)
    #     add_paragraph(doc, paragraph)

    # Save the document
    return doc



#st.title("Projects")
st.subheader( f"RFP Document Formating",divider='rainbow')
Formating = """ 
You are an expert in creating RFP response templates in word and PPT.
Search the given RFP document in entirety and  look out for specific instructions or guidelines mentioned in the document regarding formatting, layout, or content organisation.
After extracting this information, provide the following information regarding the response template formatting and layout:
Optimal document formatting and layout.
The page limit and formatting requirements for the response document as listed in the RFP document
The font size as listed in the RFP document with exception for charts, tables, and graphics
The list of items that are excluded from the page count, font, margin, and spacing requirements
Provide detailed Instructions for submission for each section or volumes in the RFP response
List out the all important Evaluation criterion for RFP response
 
"""
st.write(st.session_state.get('Formating' ,None))
st.subheader( "Do You Want To Create RFP Response Template ",divider='rainbow')
if st.button('Generate Template'):
    doc = create_document()
            #doc.save("formatted_document.docx")
    bio = io.BytesIO()
    doc.save(bio)
    st.download_button(
                label="Download Template",
                data = bio.getvalue(),
                file_name = "Template.docx",
                mime="docx"
            )

