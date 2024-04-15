import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
import ast
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING

from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
#from dotenv import load_dotenv

#load_dotenv()
#os.getenv("GOOGLE_API_KEY")
#genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

os.environ['GOOGLE_API_KEY']="AIzaSyA8TesjTLzuqovktzzhu4Vvdo3F0Tbcpvc"

# Document Creation 
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING

from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn

def set_spacing(paragraph, line_spacing):
    """
    Set line spacing for a paragraph.
    """
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = line_spacing

def set_margins(document, top, bottom, left, right):
    """
    Set margins for the entire document.
    """
    sections = document.sections
    for section in sections:
        section.top_margin = top
        section.bottom_margin = bottom
        section.left_margin = left
        section.right_margin = right


def set_heading_style(heading, font_name, font_size):
    run = heading.runs[0]
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.color.rgb = RGBColor(0, 255, 0)
    
def add_page_numbers(document):
    """
    Add page numbers to all pages in the document.
    """
    footer = document.sections[0].footer
    footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer.paragraphs[0].add_run("Page ").bold = True
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE \\* MERGEFORMAT')
    footer.paragraphs[0]._element.append(fldSimple)
    footer.paragraphs[0].add_run(" of ").bold = True
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'NUMPAGES \\* MERGEFORMAT')
    footer.paragraphs[0]._element.append(fldSimple)

def create_document():
    user_question = """ Using the given text, extract the following information and store it in a JSON format:

"Font Name": <string>
"Font Size": <number>
"Page Format": <string>
"Margins": <number> (e.g., "1")


"Page Count": <number>
"Heading Levels": <list of strings> (e.g., ["Heading 1", "Heading 2", ...])





if in case any Value is Not Prsent in the Document so just Specify Value not provided don't give me False Answer """

    embeddings = GoogleGenerativeAIEmbeddings(model = "models/embedding-001")
    new_db = FAISS.load_local("faiss_index", embeddings)
    docs = new_db.similarity_search(user_question)
    chain = get_conversational_chain()

    
    response = chain(
        {"input_documents":docs, "question": user_question}
        , return_only_outputs=True)
    
    print(type(response))
    #lines = response.strip().split('\n')
    #dictionary = {}

    #for line in lines:
        #key, value = line.split(': ')
    # Converting the value part to appropriate types
        ##try:
            #value = ast.literal_eval(value)
        #except (ValueError, SyntaxError):
            #value = value.strip('"')
        #dictionary[key.strip('"')] = value

    """
    Create a new document with specified formatting.
    """
    document = Document()
    # Add cover page
    cover_page = document.add_paragraph("Cover Page", style='Title')
    cover_page.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_page_break()

    # Set margins
    set_margins(document, Pt(54), Pt(54), Pt(54), Pt(54))
    #set_heading_style(document)
    title = document.add_heading("Introduction",1)
    set_heading_style(title,"Monotype Corsiva",18)
    paragraphs = [
        """This document provides sample text and technical resources for state, local, tribal, territorial, or 
    regional agencies as they develop Request for Information (RFI) and Request for Proposal (RFP)
    documents to procure Land Mobile Radio (LMR) subscriber units. Given that agencies vary in size,
    scope, and resources, the sample text is not intended to introduce "one-size-fits-all" approaches to
    RFP, but rather to provide examples of language common to LMR subscriber units RFIs and RFPs.
    Agencies and entities are encouraged to use the following sample texts as a roadmap to RFP
    development, modified appropriately to fit the needs of the agency or entity""" ,
        
        """In addition to the sample language, this document also includes a list of general resources that may
be helpful for users when developing RFIs or RFPs. Users are encouraged to explore these resources
to enhance their understanding of RFP development."""
    ]
    page_count = 2  # Cover page + 2 paragraphs

    for paragraph_text in paragraphs:
        if page_count >= 10:
            break

        paragraph = document.add_paragraph(paragraph_text)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        #set_spacing(paragraph,WD_LINE_SPACING.EXACTLY)

        # Set font size and font
        for run in paragraph.runs:
            run.font.size = Pt(12) 
            run.font.name = 'Times New Roman' 

        # Increment page count based on approximate lines per paragraph
        lines = paragraph_text.count('\n') + 1
        page_count += lines / 30  # Assuming 30 lines per page approximately

    # Add page numbers
    add_page_numbers(document)

    return document


##################################################################################################################################################
def get_pdf_text(pdf_docs):
    text=""
    for pdf in pdf_docs:
        pdf_reader= PdfReader(pdf)
        for page in pdf_reader.pages:
            text+= page.extract_text()
    return  text



def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = text_splitter.split_text(text)
    return chunks


def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model = "models/embedding-001")
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")


def get_conversational_chain():

    prompt_template = """
    You are an AI assistant with advanced natural language processing and information extraction capabilities. Your task is to extract key details from a given RFP document and provide a structured summary of the most important information.
Present the extracted information in a clear and concise manner, organizing it into well-structured sections. Use markdown formatting for any code snippets or technical details. Ensure that the response provides a comprehensive overview of the RFP's key details to help the user effectively respond to the request.
Do not infer any data based on previous training, strictly use only source text given below as input.If you don't know the answer, just say that you don't know, don't try to make up an answer.\n\n
    Context:\n {context}?\n
    Question: \n{question}\n

    Answer:
    """

    model = ChatGoogleGenerativeAI(model="gemini-pro",
                             temperature=0.2)

    prompt = PromptTemplate(template = prompt_template, input_variables = ["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

    return chain



def user_input(user_question):
    embeddings = GoogleGenerativeAIEmbeddings(model = "models/embedding-001")
    
    new_db = FAISS.load_local("faiss_index", embeddings)
    docs = new_db.similarity_search(user_question)

    chain = get_conversational_chain()

    
    response = chain(
        {"input_documents":docs, "question": user_question}
        , return_only_outputs=True)

    print(response)
    st.write("Reply: ", response["output_text"])




def main():
    st.set_page_config(page_title="Automating RFP Analysis",page_icon="robat:")
    st.markdown(
        """
        <div style="text-align:center;">
        <h1>Streamlining RFP Analysis</h1>
        <h3>Automated Feature Extraction & Document Template Generation</h3>
        <p> Our project automates the process of analyzing Request for Proposals (RFPs) by
          extracting essential features and generating document templates. 
          By leveraging advanced models, we streamline the tedious task of identifying key elements from RFPs, 
          enabling efficient proposal preparation. ✨</p>
        </div>

    """,
    unsafe_allow_html=True,
    )

    option = st.selectbox("Select the Options",("None","Summary of the RFP doc","Deadlines,Timelines,Contact Information","Scope of RFP","RFP Requirement","Sections to be included in RFP Response","Instruction for Submision including Formating"))
    st.subheader( f"{option}",divider='rainbow')

    if option == "Summary of the RFP doc":
        Summary= """ Summarize the key points and main content of the provided RFP document in a concise and informative manner. Highlight the following key elements:

1. Background and introduction to the project or requirement
2. Specific objectives and scope of work
3. Desired qualifications and experience of the vendor/contractor
4. Submission requirements and evaluation criteria
5. Timeline and important deadlines
Give the headline Summary Of RFP Document 
Structure your summary in a clear and organized way, using bullet points or short paragraphs to convey the most salient information. Focus on extracting the essential details and main takeaways that a reader would need to understand the RFP and determine if their company is a good fit to respond.

Avoid extraneous details or verbatim copying from the original document. Instead, synthesize the information into a high-level overview that captures the core intent and key requirements of the RFP. The summary should be approximately 300-500 words in length."""

        user_input(Summary)

    if option == "Deadlines,Timelines,Contact Information":
        Deadlines = """  
Please provide the following information from the given RFP documentation:

1. Deadlines:
   - List all key deadlines mentioned in the RFP, including the deadline for proposal submission, Submision of solicitation ,contract award, project start, and any other important milestones.
   - For each deadline, provide the specific date or timeline (e.g., "Proposal submission deadline: June 15, 2023").

2. Timelines:
   - Outline the overall project timeline as described in the RFP, including the anticipated duration of the project, key phases or milestones, and any specific timelines for deliverables or project stages.

3. Contact Information:
   - Identify the primary points of contact for the RFP, including the name, title, email address, and phone number of the person(s) responsible for managing the RFP process.

Key Points:
- Carefully review the RFP document to extract the requested information.
- Provide the details in a clear and concise format, using bullet points or a table if appropriate.
- Ensure that all dates, timelines, and contact information are accurate and match the RFP documentation.
- Highlight any important or time-sensitive deadlines that require special attention."""
        user_input(Deadlines)

    if option == "RFP Requirement":
        Requirement = """ 1. Requirement Extraction:
   - List all the technical, functional, and non-functional requirements specified in the RFP.
   - Organize the requirements in a clear, concise format (e.g., a table or bullet points).
   - For each requirement, include a brief description, any specific details or constraints, and the priority level (e.g., mandatory, desirable).

2. Submission Requirements:
   - Identify all the elements and information that the proposal submission must contain, such as:
     - Proposal format and structure
     - Required sections (e.g., executive summary, technical approach, project plan, team qualifications)
     - Formatting guidelines (e.g., page limits, font size, margins)
     - Submission deadline and method (e.g., electronic, physical)
     - Any supporting documents or attachments required

3. Other Important Information:
   - Summarize any other key details from the RFP that the client should be aware of, such as:
     - Project scope and objectives
     - Evaluation criteria and selection process
     - Budget or pricing guidelines
     - Timeline and milestones
     - Contract terms and conditions

Present the extracted information in a clear, organized manner, using appropriate headings, bullet points, and formatting to make it easy for the client to understand and reference."""
        user_input(Requirement)

    if option == "Scope of RFP":
        Scope = """ Headline:
Provide a concise headline that summarizes the key scope of the RFP.

Scope Extraction:
- Identify the main objectives and requirements of the RFP.
- Extract the key deliverables, services, or products that the client is seeking.
- Note any specific technologies, tools, or methodologies that the client has specified.
- Summarize any other important scope-related details from the RFP document.
Give the headline

Please present the extracted information in a clear and structured format."""
        user_input(Scope)
    if option == "Sections to be included in RFP Response":
        Sections = """ You are an expert in crafting RFP documents. I need your help in extracting the relevant sections that should be included in an RFP for a new project. The project is to develop a customer relationship management (CRM) system for a medium-sized retail company.

Please provide the following:
Provide a bulleted list of the key sections that the proposal should address (e.g., Executive Summary, Company Overview, Proposed Solution, Implementation Plan, Pricing, References).
   - also Provide a concise 1-2 paragraph summary of the each of the sections which are to be responded in RFP."""
        user_input(Sections)
    if option == "Instruction for Submision including Formating":
        Formating = """ 
Any specific instructions or guidelines mentioned in the document regarding formatting, layout, or content organization.
After extracting this information, suggest the following:

Optimal document formatting and layout.
Strategies for improving content structure and readability.
Recommendations for including or enhancing special elements like graphics, tables, and charts.
The page limit and formatting requirements for the Technical Volume
The font size exception for charts, tables, and graphics
The list of items that are excluded from the page count, font, margin, and spacing requirements"""
        user_input(Formating)
        st.subheader( "Do You Want To Create RFP Document Template ",divider='rainbow')
        if st.button('Yes'):
            doc = create_document()
            doc.save("formatted_document.docx")
        



        
    with st.sidebar:
        st.title("Genexa 💼")
        st.subheader("Upload your Documents")
        pdf_docs = st.file_uploader("Upload your PDF Files and Click on the Submit & Process Button", accept_multiple_files=True)
        if st.button("Submit & Process"):
            with st.spinner("Processing..."):
                raw_text = get_pdf_text(pdf_docs)
                text_chunks = get_text_chunks(raw_text)
                get_vector_store(text_chunks)
                st.success("Done")


if __name__ == "__main__":
    main()
    
